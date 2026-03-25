"""
Script para gerar relatório técnico-executivo do Regeneration Credit AI Assistant
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_heading_custom(doc, text, level=1):
    """Adiciona título customizado"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.style.font.color.rgb = RGBColor(76, 175, 80)  # Verde
    return heading

def add_paragraph_formatted(doc, text, bold=False, italic=False, size=11):
    """Adiciona parágrafo formatado"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return para

def add_code_block(doc, text):
    """Adiciona bloco de código/diagrama"""
    para = doc.add_paragraph(text)
    para.style.font.name = 'Courier New'
    para.style.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.5)
    return para

def create_report():
    """Cria o relatório completo"""
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ========================================================================
    # PÁGINA 1: CAPA E INTRODUÇÃO
    # ========================================================================
    
    # Capa
    title = doc.add_heading('Regeneration Credit AI Assistant', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Sistema de Chatbot com RAG para Projeto de Regeneração da Natureza')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    # Informações do documento
    doc.add_paragraph()
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Data: {datetime.now().strftime("%d/%m/%Y")}\n').font.size = Pt(11)
    info.add_run('Versão: 1.0.0-beta\n').font.size = Pt(11)
    info.add_run('Relatório Técnico-Executivo').font.size = Pt(11)
    
    doc.add_page_break()
    
    # 1.1 Introdução
    add_heading_custom(doc, '1. Introdução', level=1)
    
    add_paragraph_formatted(doc, 
        'O Regeneration Credit é um sistema peer-to-peer de regeneração da natureza baseado em '
        'blockchain, que conecta pessoas interessadas em apoiar projetos de impacto ambiental com '
        'produtores e regeneradores que executam ações práticas de preservação e restauração.'
    )
    
    add_paragraph_formatted(doc, 
        'Para facilitar o entendimento e a adoção do projeto por diferentes públicos, foi desenvolvido '
        'o Regeneration Credit AI Assistant, um chatbot inteligente que utiliza Inteligência Artificial '
        'avançada para responder dúvidas sobre o projeto em linguagem natural e acessível.'
    )
    
    add_heading_custom(doc, '1.1 Objetivo do Chatbot AI', level=2)
    
    add_paragraph_formatted(doc, 
        'O assistente foi projetado para democratizar o acesso à informação sobre o Regeneration Credit, '
        'permitindo que qualquer pessoa, independente do nível técnico, possa:'
    )
    
    doc.add_paragraph('Entender os conceitos fundamentais do projeto', style='List Bullet')
    doc.add_paragraph('Consultar informações sobre tokenomics e distribuição', style='List Bullet')
    doc.add_paragraph('Aprender sobre contratos inteligentes e implementação técnica', style='List Bullet')
    doc.add_paragraph('Obter orientações sobre uso do aplicativo Core RC', style='List Bullet')
    doc.add_paragraph('Esclarecer dúvidas sobre blockchain Sintrop e infraestrutura', style='List Bullet')
    
    add_heading_custom(doc, '1.2 Escopo do Relatório', level=2)
    
    add_paragraph_formatted(doc, 
        'Este documento apresenta uma visão executiva do sistema desenvolvido, abordando as tecnologias '
        'adotadas, a arquitetura da solução, o sistema RAG (Retrieval-Augmented Generation) implementado '
        'e os próximos passos planejados para evolução do projeto.'
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # PÁGINA 2: LINGUAGEM E STACK TECNOLÓGICO
    # ========================================================================
    
    add_heading_custom(doc, '2. Linguagem de Programação e Stack Tecnológico', level=1)
    
    add_heading_custom(doc, '2.1 Linguagem Principal: Python 3.10+', level=2)
    
    add_paragraph_formatted(doc, 
        'O projeto foi desenvolvido integralmente em Python 3.10+, uma escolha estratégica motivada '
        'principalmente pela forte aderência desta linguagem a projetos envolvendo Large Language Models '
        '(LLMs) e Inteligência Artificial.'
    )
    
    add_heading_custom(doc, '2.2 Justificativa da Escolha', level=2)
    
    add_paragraph_formatted(doc, 
        'Python se consolidou como a linguagem predominante no ecossistema de IA e Machine Learning, '
        'oferecendo vantagens decisivas:'
    )
    
    doc.add_paragraph(
        'Ecossistema Rico em IA/ML: Bibliotecas maduras e especializadas (LangChain, Transformers, '
        'ChromaDB) com suporte nativo e documentação abrangente',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Integração Facilitada com LLMs: APIs oficiais dos principais provedores (Anthropic, OpenAI) '
        'mantêm SDKs Python como prioridade',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Prototipagem Rápida: Sintaxe clara e bibliotecas de alto nível permitem desenvolvimento ágil, '
        'ideal para projetos de POC (Proof of Concept)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Comunidade Ativa: Vasta comunidade focada em IA com recursos, exemplos e suporte contínuo',
        style='List Bullet'
    )
    
    add_heading_custom(doc, '2.3 Stack Tecnológico Detalhado', level=2)
    
    add_paragraph_formatted(doc, 'Framework de Orquestração:', bold=True)
    add_paragraph_formatted(doc, 
        'LangChain 1.0+ - Framework especializado em construção de aplicações com LLMs, oferecendo '
        'abstrações para agentes, memória, ferramentas e integração com vector stores.'
    )
    
    add_paragraph_formatted(doc, 'Modelo de IA:', bold=True)
    add_paragraph_formatted(doc, 
        'Claude Haiku 4.5 (Anthropic) - LLM de última geração, otimizado para respostas rápidas e '
        'custo-efetivas, com forte capacidade de raciocínio e seguimento de instruções complexas.'
    )
    
    add_paragraph_formatted(doc, 'Vector Database:', bold=True)
    add_paragraph_formatted(doc, 
        'ChromaDB - Banco de dados vetorial leve e eficiente, otimizado para buscas de similaridade '
        'semântica em grandes volumes de documentos.'
    )
    
    add_paragraph_formatted(doc, 'Sistema de Embeddings:', bold=True)
    add_paragraph_formatted(doc, 
        'Sentence Transformers (all-MiniLM-L6-v2) - Modelo de embeddings multilíngue que converte '
        'texto em vetores numéricos, possibilitando buscas semânticas precisas.'
    )
    
    add_paragraph_formatted(doc, 'Interface de Usuário:', bold=True)
    add_paragraph_formatted(doc, 
        'Streamlit - Framework Python para criação rápida de interfaces web interativas. Escolhido pela '
        'facilidade de integração com Python e adequação para prova de conceito, permitindo desenvolvimento '
        'ágil sem necessidade de conhecimento profundo em front-end.'
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # PÁGINA 3: ARQUITETURA DO SISTEMA
    # ========================================================================
    
    add_heading_custom(doc, '3. Arquitetura do Sistema', level=1)
    
    add_heading_custom(doc, '3.1 Visão Geral da Arquitetura', level=2)
    
    add_paragraph_formatted(doc, 
        'O sistema foi projetado seguindo uma arquitetura modular em camadas, separando responsabilidades '
        'e facilitando manutenção e evolução:'
    )
    
    doc.add_paragraph()
    
    # Diagrama ASCII
    add_code_block(doc, '''
    ┌─────────────────────────────────────────────────────────────┐
    │                   CAMADA DE INTERFACE                       │
    │                    (Streamlit UI)                           │
    │  ┌─────────┐  ┌──────────┐  ┌─────────┐  ┌──────────┐    │
    │  │  Chat   │  │ Prompts  │  │Retriever│  │ Tokens & │    │
    │  │   Tab   │  │   Tab    │  │  Debug  │  │  Custos  │    │
    │  └─────────┘  └──────────┘  └─────────┘  └──────────┘    │
    └────────────────────┬────────────────────────────────────────┘
                         │
    ┌────────────────────┴────────────────────────────────────────┐
    │              CAMADA DE AGENTE (ReAct Loop)                  │
    │        RegenerationCreditAgent + TokensTracker              │
    │                                                              │
    │  ┌──────────────┐        ┌─────────────────┐               │
    │  │System Prompt │───────▶│  Claude Haiku   │               │
    │  │ + Memory     │        │      4.5        │               │
    │  └──────────────┘        └─────────────────┘               │
    └────────────────────┬────────────────────────────────────────┘
                         │
    ┌────────────────────┴────────────────────────────────────────┐
    │                   CAMADA RAG                                │
    │                                                              │
    │  ┌───────────────┐  ┌─────────────┐  ┌──────────────┐     │
    │  │search_general │  │search_white │  │search_contract│     │
    │  │               │  │  paper      │  │              │     │
    │  └───────┬───────┘  └──────┬──────┘  └──────┬───────┘     │
    │          │                  │                │              │
    │          └──────────────────┼────────────────┘              │
    │                             ▼                               │
    │              ┌─────────────────────────────┐                │
    │              │    Vector Store Manager     │                │
    │              │        (ChromaDB)           │                │
    │              └─────────────────────────────┘                │
    └────────────────────┬────────────────────────────────────────┘
                         │
    ┌────────────────────┴────────────────────────────────────────┐
    │                   CAMADA DE DADOS                           │
    │                                                              │
    │  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────┐   │
    │  │Contratos │  │   Docs   │  │Whitepaper│  │ Manuais  │   │
    │  │Solidity  │  │Markdown  │  │    RC    │  │  e Guias │   │
    │  │(57 arqs) │  │(66 arqs) │  │          │  │          │   │
    │  └──────────┘  └──────────┘  └──────────┘  └──────────┘   │
    └─────────────────────────────────────────────────────────────┘
    ''')
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '3.2 Descrição dos Componentes', level=2)
    
    add_paragraph_formatted(doc, 'Camada de Interface (Streamlit):', bold=True)
    add_paragraph_formatted(doc, 
        'Responsável pela interação com o usuário através de múltiplas abas especializadas. Além do '
        'chat principal, oferece transparência através de abas de debug (Retriever, Tokens & Custos) '
        'e inspeção (Prompts).'
    )
    
    add_paragraph_formatted(doc, 'Camada de Agente (RegenerationCreditAgent):', bold=True)
    add_paragraph_formatted(doc, 
        'Núcleo inteligente do sistema. Implementa um loop ReAct (Reasoning + Acting) manual, onde '
        'o agente raciocina sobre a pergunta do usuário, decide quais ferramentas usar, executa buscas '
        'e formula respostas. Inclui sistema de memória conversacional para manter contexto entre turnos.'
    )
    
    add_paragraph_formatted(doc, 'Camada RAG (Retrieval-Augmented Generation):', bold=True)
    add_paragraph_formatted(doc, 
        'Gerencia o acesso ao conhecimento através de 4 ferramentas especializadas que buscam informações '
        'no vector store. Realiza auditoria completa de cada busca (tempo, scores, metadados) para '
        'debugging e otimização.'
    )
    
    add_paragraph_formatted(doc, 'Camada de Dados:', bold=True)
    add_paragraph_formatted(doc, 
        'Repositório de conhecimento com 9 tipos de fontes diferentes, totalizando mais de 120 documentos '
        'indexados e pesquisáveis semanticamente.'
    )
    
    doc.add_page_break()
    
    add_heading_custom(doc, '3.3 Fluxo de Processamento de Pergunta', level=2)
    
    add_code_block(doc, '''
    USUÁRIO                                             SISTEMA
       │
       │  "O que é Regeneration Credit?"
       ├───────────────────────────────────────────────▶  1. RECEBE PERGUNTA
       │                                                   │
       │                                                   ▼
       │                                          2. AGENTE ANALISA
       │                                             (Claude Haiku)
       │                                                   │
       │                                                   ▼
       │                                       3. DECIDE FERRAMENTAS
       │                                          (search_whitepaper)
       │                                                   │
       │                                                   ▼
       │                                      4. BUSCA NO VECTOR STORE
       │                                         (Embeddings + ChromaDB)
       │                                                   │
       │                                                   ▼
       │                                      5. RECUPERA DOCUMENTOS
       │                                         (Top-5 mais relevantes)
       │                                                   │
       │                                                   ▼
       │                                      6. AGENTE FORMULA RESPOSTA
       │                                         (Com contexto dos docs)
       │                                                   │
       │  Resposta completa                                ▼
       │◀─────────────────────────────────────────  7. RETORNA RESPOSTA
       │                                              + Tracking (tokens, tempo)
       │
    ''')
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '3.4 Padrões de Design Utilizados', level=2)
    
    add_paragraph_formatted(doc, 'Loop ReAct Manual:', bold=True)
    add_paragraph_formatted(doc, 
        'Implementação customizada do padrão ReAct (Reason + Act), permitindo controle granular sobre '
        'cada iteração do agente, incluindo tracking detalhado de tokens e custos.'
    )
    
    add_paragraph_formatted(doc, 'Separação de Responsabilidades:', bold=True)
    add_paragraph_formatted(doc, 
        'Cada camada possui responsabilidades bem definidas: interface (UX), agente (lógica), RAG (retrieval), '
        'dados (armazenamento). Facilita testes, manutenção e evolução independente.'
    )
    
    add_paragraph_formatted(doc, 'Modularização:', bold=True)
    add_paragraph_formatted(doc, 
        'Componentes independentes e reutilizáveis (TokensTracker, VectorStoreManager, RAGTools) que podem '
        'ser testados isoladamente e utilizados em outros contextos.'
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # PÁGINA 4: SISTEMA RAG
    # ========================================================================
    
    add_heading_custom(doc, '4. Sistema RAG Estruturado', level=1)
    
    add_heading_custom(doc, '4.1 Conceito e Funcionamento do RAG', level=2)
    
    add_paragraph_formatted(doc, 
        'RAG (Retrieval-Augmented Generation) é uma técnica que combina a capacidade de geração de texto '
        'de LLMs com a recuperação de informações relevantes de uma base de conhecimento específica.'
    )
    
    add_paragraph_formatted(doc, 
        'Ao invés de depender apenas do conhecimento pré-treinado do modelo de IA (que pode estar '
        'desatualizado ou não incluir informações específicas do projeto), o sistema:'
    )
    
    doc.add_paragraph(
        'Busca documentos relevantes na base de conhecimento do projeto',
        style='List Number'
    )
    doc.add_paragraph(
        'Fornece esses documentos como contexto para o LLM',
        style='List Number'
    )
    doc.add_paragraph(
        'O LLM gera respostas baseadas nesse contexto específico',
        style='List Number'
    )
    
    add_paragraph_formatted(doc, 'Vantagens para o Projeto:', bold=True)
    
    doc.add_paragraph(
        'Respostas Precisas: Baseadas em documentação oficial do projeto',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Sempre Atualizado: Basta atualizar os documentos e reprocessar',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Rastreável: Cada resposta pode ser vinculada às fontes consultadas',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Econômico: Reduz necessidade de fine-tuning de modelos',
        style='List Bullet'
    )
    
    add_heading_custom(doc, '4.2 Fontes de Dados (9 tipos, 120+ documentos)', level=2)
    
    add_paragraph_formatted(doc, 
        'O sistema indexa e processa 9 tipos diferentes de fontes de conhecimento:'
    )
    
    doc.add_paragraph()
    add_paragraph_formatted(doc, '1. Contratos Solidity (57 arquivos)', bold=True)
    add_paragraph_formatted(doc, 
        'Implementação técnica dos smart contracts: Pools, Rules, interfaces e tipos.'
    )
    
    add_paragraph_formatted(doc, '2. Documentação Markdown (66 arquivos)', bold=True)
    add_paragraph_formatted(doc, 
        'Documentação técnica gerada automaticamente do código fonte.'
    )
    
    add_paragraph_formatted(doc, '3. README.md', bold=True)
    add_paragraph_formatted(doc, 
        'Visão geral do projeto principal, arquitetura e instruções.'
    )
    
    add_paragraph_formatted(doc, '4. CHANGELOG.md', bold=True)
    add_paragraph_formatted(doc, 
        'Histórico completo de mudanças e evolução do projeto.'
    )
    
    add_paragraph_formatted(doc, '5. Whitepaper Regeneration Credit', bold=True)
    add_paragraph_formatted(doc, 
        'Visão estratégica, tokenomics e regras de negócio do sistema.'
    )
    
    add_paragraph_formatted(doc, '6. Manual Core RC', bold=True)
    add_paragraph_formatted(doc, 
        'Guia completo do usuário do aplicativo Core (cadastro, níveis, saques, certificados).'
    )
    
    add_paragraph_formatted(doc, '7. Tutorial de Carteira MetaMask', bold=True)
    add_paragraph_formatted(doc, 
        'Passo a passo para criar e configurar carteira blockchain.'
    )
    
    add_paragraph_formatted(doc, '8. Guia de Mineração Sintrop', bold=True)
    add_paragraph_formatted(doc, 
        'Instruções técnicas para setup de nós e mineração na blockchain.'
    )
    
    add_paragraph_formatted(doc, '9. Whitepaper Sintrop Blockchain', bold=True)
    add_paragraph_formatted(doc, 
        'Arquitetura da blockchain, consenso Proof-of-Work e especificações técnicas.'
    )
    
    doc.add_page_break()
    
    add_heading_custom(doc, '4.3 Ferramentas RAG Especializadas', level=2)
    
    add_paragraph_formatted(doc, 
        'Para otimizar a busca e reduzir ruído, foram criadas 4 ferramentas especializadas que o agente '
        'pode escolher conforme o tipo de pergunta:'
    )
    
    doc.add_paragraph()
    
    add_code_block(doc, '''
    ┌─────────────────────────────────────────────────────────────┐
    │                    FERRAMENTAS RAG                          │
    └─────────────────────────────────────────────────────────────┘
    
    1. search_general
       │
       ├─▶ Escopo: Busca ampla em TODAS as fontes
       ├─▶ Uso: Perguntas gerais, documentação, manuais, tutoriais
       └─▶ Exemplos: "Como usar o app Core?", "Como criar carteira?"
    
    2. search_whitepaper
       │
       ├─▶ Escopo: Apenas Whitepaper RC
       ├─▶ Uso: Visão, propósito, tokenomics, regras de negócio
       └─▶ Exemplos: "O que é RC?", "Como funciona distribuição?"
    
    3. search_contracts
       │
       ├─▶ Escopo: Apenas contratos Solidity
       ├─▶ Uso: Implementação técnica, funções, variáveis
       └─▶ Exemplos: "Como funciona o Pool?", "Quais funções do Rules?"
    
    4. consult_tokenomics_guide
       │
       ├─▶ Escopo: Documento de síntese completo
       ├─▶ Uso: Cálculos, fórmulas, valores de referência
       └─▶ Exemplos: "Calcule meu score", "Quais parâmetros usar?"
    ''')
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '4.4 Processamento e Indexação', level=2)
    
    add_paragraph_formatted(doc, 
        'O pipeline de processamento garante que documentos de diferentes formatos sejam '
        'uniformemente indexados:'
    )
    
    doc.add_paragraph()
    
    add_code_block(doc, '''
    DOCUMENTOS FONTE          PROCESSAMENTO           VECTOR STORE
    
    ┌─────────────┐
    │  PDFs       │───┐
    └─────────────┘   │
                      ├──▶  1. Conversão         ┌─────────────┐
    ┌─────────────┐   │     (Docling)           │             │
    │  Markdown   │───┤                      ───▶│  ChromaDB   │
    └─────────────┘   │     2. Chunking          │             │
                      │     (1000 chars)         │  Embeddings │
    ┌─────────────┐   │     (200 overlap)        │  Vetoriais  │
    │  Solidity   │───┤                          │             │
    └─────────────┘   │     3. Embeddings    ───▶│  Busca      │
                      │     (all-MiniLM)         │  Semântica  │
    ┌─────────────┐   │                          │             │
    │  README     │───┘     4. Indexação         └─────────────┘
    └─────────────┘         (Metadados)
    ''')
    
    doc.add_paragraph()
    
    add_paragraph_formatted(doc, 'Estratégia de Chunking:', bold=True)
    add_paragraph_formatted(doc, 
        'Documentos são divididos em chunks (pedaços) de 1000 caracteres com sobreposição de 200 '
        'caracteres. Isso garante que informações contextuais não sejam perdidas nas divisões e permite '
        'recuperação mais precisa.'
    )
    
    add_paragraph_formatted(doc, 'Metadados Enriquecidos:', bold=True)
    add_paragraph_formatted(doc, 
        'Cada chunk é enriquecido com metadados (source_type, source, título) que permitem filtragem '
        'e rastreamento da origem das informações.'
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # PÁGINA 5: TECNOLOGIAS E FERRAMENTAS
    # ========================================================================
    
    add_heading_custom(doc, '5. Tecnologias e Ferramentas', level=1)
    
    add_heading_custom(doc, '5.1 Frameworks e Bibliotecas Core', level=2)
    
    add_paragraph_formatted(doc, 'LangChain 1.0+ (Orquestração de IA):', bold=True)
    add_paragraph_formatted(doc, 
        'Framework Python especializado em desenvolvimento de aplicações com LLMs. Fornece abstrações '
        'de alto nível para agentes, memória conversacional, ferramentas e integração com vector stores. '
        'Escolhido por ser o padrão da indústria e possuir ampla documentação e comunidade ativa.'
    )
    
    add_paragraph_formatted(doc, 'Claude Haiku 4.5 - Anthropic (LLM):', bold=True)
    add_paragraph_formatted(doc, 
        'Modelo de linguagem de última geração da Anthropic, otimizado para velocidade e custo-efetividade. '
        'Oferece capacidades avançadas de raciocínio, seguimento preciso de instruções complexas e '
        'comportamento mais previsível comparado a modelos concorrentes. Ideal para aplicações de produção '
        'que exigem respostas rápidas e consistentes.'
    )
    
    add_paragraph_formatted(doc, 'ChromaDB (Vector Database):', bold=True)
    add_paragraph_formatted(doc, 
        'Banco de dados vetorial open-source, leve e fácil de usar. Projetado especificamente para '
        'aplicações de IA, permite armazenamento e busca eficiente de embeddings. Escolhido por sua '
        'simplicidade de configuração, adequação para POC e capacidade de escalar para produção.'
    )
    
    add_paragraph_formatted(doc, 'HuggingFace Sentence Transformers (Embeddings):', bold=True)
    add_paragraph_formatted(doc, 
        'Biblioteca para geração de embeddings de texto usando modelos transformer. O modelo all-MiniLM-L6-v2 '
        'foi escolhido por oferecer excelente equilíbrio entre qualidade, velocidade e tamanho, com suporte '
        'multilíngue (incluindo português).'
    )
    
    add_heading_custom(doc, '5.2 Interface e Experiência do Usuário', level=2)
    
    add_paragraph_formatted(doc, 'Streamlit (Framework de UI):', bold=True)
    add_paragraph_formatted(doc, 
        'Framework Python para criação rápida de aplicações web interativas. Destacam-se as vantagens:'
    )
    
    doc.add_paragraph(
        'Integração Nativa com Python: Sem necessidade de JavaScript ou frameworks front-end complexos',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Desenvolvimento Ágil: Ideal para POC e MVPs, permitindo iteração rápida',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Componentes Prontos: Widgets interativos (chat, tabs, forms) sem configuração adicional',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Deploy Facilitado: Streamlit Cloud e outras plataformas oferecem deploy com poucos cliques',
        style='List Bullet'
    )
    
    add_paragraph_formatted(doc, 'Pandas (Exportação de Dados):', bold=True)
    add_paragraph_formatted(doc, 
        'Biblioteca de análise de dados utilizada para exportação de métricas e conversas em formato CSV, '
        'permitindo análises posteriores em ferramentas como Excel.'
    )
    
    add_heading_custom(doc, '5.3 Tracking e Observabilidade', level=2)
    
    add_paragraph_formatted(doc, 
        'Para garantir transparência, debugging eficiente e otimização de custos, o sistema implementa '
        'múltiplas camadas de tracking:'
    )
    
    add_paragraph_formatted(doc, 'Tokens Tracker:', bold=True)
    add_paragraph_formatted(doc, 
        'Componente customizado que contabiliza precisamente tokens consumidos em cada chamada ao LLM, '
        'distinguindo entre tokens de input, output, cache creation e cache read. Permite análise detalhada '
        'de uso por componente.'
    )
    
    add_paragraph_formatted(doc, 'Pricing Calculator:', bold=True)
    add_paragraph_formatted(doc, 
        'Calcula custos em tempo real baseado nas tabelas de preços da Anthropic, fornecendo estimativas '
        'precisas de custo por turno, sessão e componente. Essencial para controle de orçamento e ROI.'
    )
    
    add_paragraph_formatted(doc, 'Retriever Audits:', bold=True)
    add_paragraph_formatted(doc, 
        'Sistema de auditoria que registra cada busca no vector store com metadados completos: query enviada, '
        'documentos retornados, scores de similaridade, tempo de execução e fonte dos documentos. Permite '
        'debugging do sistema RAG e identificação de melhorias.'
    )
    
    add_paragraph_formatted(doc, 'LangSmith (Opcional):', bold=True)
    add_paragraph_formatted(doc, 
        'Plataforma de observabilidade da LangChain que oferece rastreamento end-to-end de todas as '
        'interações, incluindo traces de execução, latências e erros. Útil para análise profunda e debugging '
        'em produção.'
    )
    
    doc.add_page_break()
    
    add_heading_custom(doc, '5.4 Processamento de Documentos', level=2)
    
    add_paragraph_formatted(doc, 'Docling (Conversão de PDFs):', bold=True)
    add_paragraph_formatted(doc, 
        'Ferramenta avançada de conversão de PDFs para Markdown, preservando estrutura, formatação e '
        'metadados. Essencial para processar whitepapers e manuais em PDF mantendo qualidade.'
    )
    
    add_paragraph_formatted(doc, 'LangChain Document Loaders:', bold=True)
    add_paragraph_formatted(doc, 
        'Conjunto de loaders especializados para diferentes formatos de documentos (Markdown, Python, '
        'Solidity). Normaliza documentos de diferentes fontes em estrutura uniforme.'
    )
    
    add_paragraph_formatted(doc, 'LangChain Text Splitters:', bold=True)
    add_paragraph_formatted(doc, 
        'Componentes inteligentes de divisão de texto que respeitam limites semânticos (parágrafos, seções) '
        'ao invés de simplesmente dividir por caracteres, melhorando qualidade do retrieval.'
    )
    
    doc.add_page_break()
    
    # ========================================================================
    # PÁGINA 6: PRÓXIMOS PASSOS
    # ========================================================================
    
    add_heading_custom(doc, '6. Próximos Passos e Conclusão', level=1)
    
    add_heading_custom(doc, '6.1 Roadmap de Desenvolvimento', level=2)
    
    add_paragraph_formatted(doc, 
        'O projeto encontra-se em versão beta funcional, com três frentes principais de evolução planejadas:'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '(i) Sistema de Log de Conversas de Usuários para Análise', level=3)
    
    add_paragraph_formatted(doc, 'Objetivo:', bold=True)
    add_paragraph_formatted(doc, 
        'Coletar e analisar conversas reais de usuários da versão beta para identificar padrões de uso, '
        'dúvidas recorrentes, qualidade das respostas e oportunidades de melhoria.'
    )
    
    add_paragraph_formatted(doc, 'Status Atual:', bold=True)
    add_paragraph_formatted(doc, 
        'O sistema já implementa salvamento automático de conversas em formato JSON enriquecido, incluindo '
        'mensagens, tokens, custos, audits de busca e metadados de sessão. Cada conversa recebe um ID único '
        'e é armazenada em data/conversations/.'
    )
    
    add_paragraph_formatted(doc, 'Melhorias Propostas:', bold=True)
    
    doc.add_paragraph(
        'Dashboard de Analytics: Visualização agregada de métricas (perguntas mais comuns, tempo médio de '
        'resposta, satisfação estimada, tópicos mais consultados)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Análise de Padrões: Identificação automática de clusters de perguntas similares e gaps de conhecimento',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Sistema de Feedback: Permitir que usuários avaliem respostas (útil/não útil) para refinamento contínuo',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Métricas de Qualidade: Análise de cobertura (perguntas respondidas vs. não respondidas), latência '
        'e acurácia baseada em feedback',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Exportação para BI: Integração com ferramentas de Business Intelligence para análises avançadas',
        style='List Bullet'
    )
    
    add_paragraph_formatted(doc, 'Impacto Esperado:', bold=True)
    add_paragraph_formatted(doc, 
        'Ciclo de melhoria contínua baseado em dados reais de uso, permitindo refinamento do system prompt, '
        'expansão da base de conhecimento e otimização das ferramentas RAG.'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '(ii) Integração com Aplicativo Regeneration Credit', level=3)
    
    add_paragraph_formatted(doc, 'Objetivo:', bold=True)
    add_paragraph_formatted(doc, 
        'Integrar o chatbot diretamente no aplicativo móvel/web do Regeneration Credit, oferecendo '
        'assistência contextual aos usuários durante sua jornada.'
    )
    
    add_paragraph_formatted(doc, 'Implementações Necessárias:', bold=True)
    
    doc.add_paragraph(
        'API REST: Desenvolvimento de endpoints RESTful para comunicação entre app e chatbot (POST /chat, '
        'GET /history, etc.)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Autenticação e Autorização: Sistema de tokens JWT para segurança, garantindo que apenas usuários '
        'autenticados acessem o chatbot',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Sincronização de Dados: Acesso a dados do perfil do usuário (nível, atividades, saldo) para '
        'personalizar respostas',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Contexto Personalizado: Respostas adaptadas ao estado do usuário (iniciante vs. avançado, tipo '
        'de usuário, histórico de atividades)',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Deep Links: Chatbot pode sugerir ações específicas no app com links diretos para telas relevantes',
        style='List Bullet'
    )
    
    add_paragraph_formatted(doc, 'Benefícios:', bold=True)
    add_paragraph_formatted(doc, 
        'Redução de fricção no onboarding de novos usuários, suporte 24/7 sem necessidade de equipe humana, '
        'aumento de engajamento e retenção através de assistência proativa.'
    )
    
    doc.add_paragraph()
    
    add_heading_custom(doc, '(iii) Aprimoramento do Sistema RAG', level=3)
    
    add_paragraph_formatted(doc, 'Objetivo:', bold=True)
    add_paragraph_formatted(doc, 
        'Melhorar a performance de recuperação (retrieve) do sistema RAG, garantindo que documentos mais '
        'relevantes sejam encontrados, enquanto mantém ou reduz o consumo de tokens (custos).'
    )
    
    add_paragraph_formatted(doc, 'Estratégias Propostas:', bold=True)
    
    add_paragraph_formatted(doc, '1. Re-ranking de Resultados:', bold=True)
    add_paragraph_formatted(doc, 
        'Implementar um segundo estágio de ranking usando modelos cross-encoder após a busca inicial. '
        'Busca vetorial retorna top-20, cross-encoder reordena e seleciona top-5 realmente mais relevantes.'
    )
    
    add_paragraph_formatted(doc, '2. Filtros Adaptativos Inteligentes:', bold=True)
    add_paragraph_formatted(doc, 
        'Sistema que analisa a pergunta e automaticamente aplica filtros de metadados (source_type) para '
        'reduzir espaço de busca. Ex: "Como minerar?" → aplica filtro para guias técnicos.'
    )
    
    add_paragraph_formatted(doc, '3. Cache de Embeddings:', bold=True)
    add_paragraph_formatted(doc, 
        'Perguntas frequentes têm embeddings pré-computados, acelerando busca e reduzindo custos de embedding.'
    )
    
    add_paragraph_formatted(doc, '4. Otimização de Chunk Size:', bold=True)
    add_paragraph_formatted(doc, 
        'Testar diferentes tamanhos de chunk (500, 1000, 1500) e estratégias de overlap para encontrar '
        'configuração ótima entre relevância e consumo de tokens.'
    )
    
    add_paragraph_formatted(doc, '5. Compressão de Contexto:', bold=True)
    add_paragraph_formatted(doc, 
        'Usar modelos de sumarização para comprimir documentos recuperados antes de enviar ao LLM, '
        'mantendo informações essenciais mas reduzindo tokens.'
    )
    
    add_paragraph_formatted(doc, '6. Embeddings Híbridos:', bold=True)
    add_paragraph_formatted(doc, 
        'Combinar busca densa (embeddings vetoriais) com busca esparsa (BM25/TF-IDF) para melhor '
        'cobertura. Especialmente útil para termos técnicos específicos.'
    )
    
    add_paragraph_formatted(doc, 'Métricas de Sucesso:', bold=True)
    add_paragraph_formatted(doc, 
        'Redução de 30% no consumo médio de tokens por consulta, aumento de 20% na precisão de retrieval '
        '(medido por feedback de usuários), redução de 40% no tempo de resposta.'
    )
    
    doc.add_paragraph()
    doc.add_page_break()
    
    add_heading_custom(doc, '6.2 Considerações Finais', level=2)
    
    add_paragraph_formatted(doc, 'Status Atual do Projeto:', bold=True)
    add_paragraph_formatted(doc, 
        'O Regeneration Credit AI Assistant encontra-se em estágio beta funcional, com todas as '
        'funcionalidades core implementadas e testadas:'
    )
    
    doc.add_paragraph(
        'Sistema RAG completo com 4 ferramentas especializadas e 120+ documentos indexados',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Interface Streamlit profissional com múltiplas abas de funcionalidade e debug',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Tracking detalhado de tokens, custos e performance de retrieval',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Salvamento automático de conversas com metadados enriquecidos',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Sistema de memória conversacional para diálogos contextualizados',
        style='List Bullet'
    )
    
    doc.add_paragraph()
    
    add_paragraph_formatted(doc, 'Benefícios Alcançados:', bold=True)
    
    add_paragraph_formatted(doc, 
        'Democratização do Acesso: Qualquer pessoa, independente do nível técnico, pode entender conceitos '
        'complexos do projeto através de explicações claras e acessíveis.'
    )
    
    add_paragraph_formatted(doc, 
        'Redução de Carga de Suporte: Dúvidas recorrentes são respondidas automaticamente 24/7, liberando '
        'equipe para questões mais complexas.'
    )
    
    add_paragraph_formatted(doc, 
        'Onboarding Acelerado: Novos usuários conseguem aprender sobre o projeto de forma autoguiada e '
        'interativa, reduzindo fricção de entrada.'
    )
    
    add_paragraph_formatted(doc, 
        'Base para Evolução: Arquitetura modular e tracking detalhado permitem iterações rápidas baseadas '
        'em dados reais de uso.'
    )
    
    doc.add_paragraph()
    
    add_paragraph_formatted(doc, 'Potencial de Expansão:', bold=True)
    
    add_paragraph_formatted(doc, 
        'O sistema foi projetado para evoluir além de um chatbot de suporte. Possibilidades futuras incluem:'
    )
    
    doc.add_paragraph(
        'Assistente de Onboarding: Guiar novos usuários passo a passo através do processo de cadastro e '
        'primeiras ações',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Calculadora Interativa: Simular cenários de tokenomics e distribuição de forma conversacional',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Suporte Multilíngue: Expandir para outros idiomas aproveitando capacidade multilíngue dos modelos',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Agente Proativo: Enviar notificações e dicas personalizadas baseadas no perfil e atividades do usuário',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Integração com Blockchain: Consultar dados on-chain em tempo real para respostas ainda mais '
        'precisas e atualizadas',
        style='List Bullet'
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_paragraph_formatted(doc, 
        'O Regeneration Credit AI Assistant representa um passo importante na democratização do acesso '
        'à informação sobre o projeto, combinando tecnologias de ponta em IA com design focado em '
        'experiência do usuário. A base sólida construída permite evolução contínua e adaptação às '
        'necessidades emergentes da comunidade Regeneration Credit.',
        italic=True
    )
    
    # Rodapé
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('—' * 40)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_text = doc.add_paragraph(
        f'Regeneration Credit AI Assistant - Relatório Técnico-Executivo v1.0\n'
        f'Gerado em: {datetime.now().strftime("%d/%m/%Y às %H:%M")}'
    )
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text.runs[0].font.size = Pt(9)
    footer_text.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    # Salvar documento
    output_path = 'Relatorio_Regeneration_Credit_AI_Assistant.docx'
    doc.save(output_path)
    print(f"Relatorio gerado com sucesso: {output_path}")
    return output_path

if __name__ == "__main__":
    create_report()

